import os
from docx import Document
from docx.shared import Inches
from PIL import Image
from settings import INPUT_DIR, OUTPUT_DIR, IMAGE_EXTENSIONS, TEXT_EXTENSION

def find_matching_pairs():
    """查找所有匹配的图片和文本文件"""
    files = os.listdir(INPUT_DIR)
    pairs = []
    
    # 构建文件名基础名到全路径的映射
    image_bases = {}
    text_bases = {}
    
    for file in files:
        base, ext = os.path.splitext(file)
        if ext.lower() in IMAGE_EXTENSIONS:
            image_bases[base] = file
        elif ext.lower() == TEXT_EXTENSION:
            text_bases[base] = file
    
    # 匹配成对文件
    for base in image_bases:
        if base in text_bases:
            pairs.append((
                os.path.join(INPUT_DIR, image_bases[base]),
                os.path.join(INPUT_DIR, text_bases[base])
            ))
    
    return pairs

def create_document(pairs):
    """生成Word文档"""
    doc = Document()
    doc.add_heading('场景提示词', level=1)
    
    for idx, (img_path, txt_path) in enumerate(pairs, 1):
        # 插入图片
        try:
            with Image.open(img_path) as img:
                doc.add_picture(img_path, width=Inches(6.0))
                doc.add_paragraph(f"场景图片: {os.path.basename(img_path)}")
        except Exception as e:
            print(f"❌ 图片插入失败 [{os.path.basename(img_path)}]: {e}")
            continue
        
        # 插入文本
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                doc.add_paragraph(f"场景提示词:\n{f.read()}")
        except Exception as e:
            print(f"❌ 文本读取失败 [{os.path.basename(txt_path)}]: {e}")
            continue
        
        doc.add_paragraph("="*40)  # 分隔线
        print(f"\r🔄 已处理 {idx}/{len(pairs)} 组文件", end="")
    
    # 保存文档
    output_path = os.path.join(OUTPUT_DIR, "output.docx")
    doc.save(output_path)
    print(f"\n✅ 文档已生成: {output_path}")

def show_welcome():
    """欢迎界面"""
    print("="*50)
    print("📝 图片文本自动合并工具")
    print("✨ 特点:")
    print("  - 支持任意文件名（中文/数字/英文）")
    print("  - 自动匹配同名图片和文本")
    print(f"  - 输入目录: {INPUT_DIR}")
    print(f"  - 输出目录: {OUTPUT_DIR}")
    print("="*50)

if __name__ == "__main__":
    show_welcome()
    pairs = find_matching_pairs()
    
    if not pairs:
        print("⚠️ 未找到匹配的图片和文本文件！请检查：")
        print(f"1. 文件是否放在 {INPUT_DIR} 目录")
        print("2. 是否有同名文件（如 张三.jpg 和 张三.txt）")
    else:
        print(f"🔍 找到 {len(pairs)} 组匹配文件")
        create_document(pairs)
    input("按回车键退出...")